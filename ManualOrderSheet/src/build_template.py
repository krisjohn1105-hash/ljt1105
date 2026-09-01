# -*- coding: utf-8 -*-
"""직접 작성한 지시서(.docx)를 자리표시자 템플릿으로 변환한다.

원본 서식(글꼴/표 너비/머리글/사용인감)을 그대로 두고 값만 {{KEY}} 로 바꾼다.
양식이 바뀌면 새 원본으로 이 스크립트를 다시 돌리면 된다.

    python -m src.build_template                # 두 종류 모두 재생성
    python -m src.build_template dividend       # 배당지급만
    python -m src.build_template odd_lot        # 단주대금만
    python -m src.build_template dividend 원본.docx
"""
from __future__ import annotations  # Python 3.9 호환 (X | None 표기)
import sys
from pathlib import Path

from docx import Document
from docx.table import _Cell

try:
    from . import config
    from .docx_util import set_paragraph_text, set_cell_text, delete_row
except ImportError:  # 스크립트로 직접 실행할 때
    sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
    from src import config
    from src.docx_util import set_paragraph_text, set_cell_text, delete_row

# 종류별 원본 지시서와 내역 표의 열 구성
SOURCES = {
    "dividend": (
        "20260825_두나미스자산운용 수기운용지시서_신한수탁_차입주식배당지급.docx",
        # 구분 | 해당 펀드 | 반영일자 | 대차수량 | 금액 | 거래유형명
        ["NO", "FUND", "APPLY_DATE", "QTY", "AMOUNT", "TRADE_DESC"],
        2,   # 합계 행에서 금액이 들어가는 tc 번호
    ),
    "odd_lot": (
        "20260901_두나미스자산운용 수기운용지시서_기업수탁_차입주식단주대금지급.docx",
        # 구분 | 해당 펀드 | 반영일자 | 지급금액 | 거래유형명
        ["NO", "FUND", "APPLY_DATE", "AMOUNT", "TRADE_DESC"],
        1,
    ),
}

ITEM_ROW = 1   # 내역 표에서 반복 행 템플릿이 되는 행


def tc(table, row_idx, tc_idx):
    return _Cell(table.rows[row_idx]._tr.tc_lst[tc_idx], table)


def build(source: Path, dest: Path, columns, total_tc: int) -> Path:
    doc = Document(str(source))
    head, items = doc.tables[0], doc.tables[1]

    # --- 머리글: 담당자 ------------------------------------------------
    for paragraph in doc.sections[0].header.paragraphs:
        if "담당:" in paragraph.text:
            text = paragraph.text
            text = text.replace(config.CONTACT_NAME, "{{CONTACT_NAME}}")
            text = text.replace(config.CONTACT_PHONE, "{{CONTACT_PHONE}}")
            set_paragraph_text(paragraph, text)

    # --- 상단 표: 문서번호 / 일자 / 수신 / 참조 / 제목 -----------------
    set_cell_text(tc(head, 1, 1), " {{DOC_NO_PREFIX}} 제{{DOC_NO}}")
    set_cell_text(tc(head, 1, 2), "{{DOC_DATE}}")
    set_cell_text(tc(head, 2, 1), " {{RECIPIENT}}")
    set_cell_text(tc(head, 3, 1), " {{CC}}")
    set_cell_text(tc(head, 4, 1), " {{SUBJECT}}")

    # --- 본문 문단 ----------------------------------------------------
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("2."):
            set_paragraph_text(
                paragraph,
                "2. 당사 금일 발생한 {{TRADE_TYPE}} 관련하여 "
                "{{ORDER_DATE_KR}} 수기운용지시를 반영 요청 드립니다.",
            )
        elif text.startswith("수취계좌:"):
            set_paragraph_text(paragraph, "수취계좌: {{RECEIVE_ACCOUNT}}")
        elif text.startswith("주식회사"):
            set_paragraph_text(paragraph, "{{COMPANY_NAME}}")
        elif text.startswith("대표이사"):
            set_paragraph_text(paragraph, "대표이사 {{CEO_NAME}} (인)")

    # --- 내역 표 ------------------------------------------------------
    # 원본에 내역 행이 여러 개면 첫 행만 남기고 지운다 (첫 행이 반복 템플릿).
    while len(items.rows) > ITEM_ROW + 2:      # 머리행 + 내역 1행 + 합계행
        delete_row(items, ITEM_ROW + 1)

    for idx, key in enumerate(columns):
        set_cell_text(tc(items, ITEM_ROW, idx), "{{%s}}" % key)
    # 거래유형명 칸은 두 줄(유형명 / 종목 상세)로 쓴다.
    set_cell_text(tc(items, ITEM_ROW, len(columns) - 1),
                  "{{TRADE_NAME}}\n{{TRADE_DETAIL}}")
    set_cell_text(tc(items, ITEM_ROW + 1, total_tc), "{{TOTAL_AMOUNT}}")

    dest.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dest))
    return dest


def main():
    args = sys.argv[1:]
    kinds = [args[0]] if args and args[0] in SOURCES else list(SOURCES)
    override = None
    if args and args[0] in SOURCES and len(args) > 1:
        override = Path(args[1])
    elif args and args[0] not in SOURCES:
        raise SystemExit(f"알 수 없는 종류: {args[0]} (dividend | odd_lot)")

    for kind in kinds:
        filename, columns, total_tc = SOURCES[kind]
        source = override or (config.TEMPLATE_DIR / filename)
        if not source.exists():
            raise SystemExit(f"원본 지시서를 찾을 수 없습니다: {source}")
        out = build(source, config.TEMPLATE_PATHS[kind], columns, total_tc)
        print(f"[{kind}] 템플릿 생성: {out.name}  (원본: {source.name})")


if __name__ == "__main__":
    main()
