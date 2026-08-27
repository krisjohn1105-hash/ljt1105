# -*- coding: utf-8 -*-
"""python-docx 보조 함수.

문단 안의 run 이 잘게 쪼개져 있으면 '{{PLACEHOLDER}}' 치환이 불가능하므로,
첫 run 의 서식을 유지한 채 문단 텍스트를 통째로 갈아끼우는 방식을 쓴다.
"""
from __future__ import annotations  # Python 3.9 호환 (X | None 표기)
import copy


def iter_paragraphs(doc):
    """본문 + 표(중첩 표 포함) + 머리글/바닥글의 모든 문단을 순회한다."""
    yield from _iter_container(doc)
    for section in doc.sections:
        parts = [section.header, section.footer]
        # 아래 두 종류는 접근만 해도 파트가 새로 생성되므로 실제 사용 중일 때만 훑는다.
        if section.different_first_page_header_footer:
            parts += [section.first_page_header, section.first_page_footer]
        if doc.settings.odd_and_even_pages_header_footer:
            parts += [section.even_page_header, section.even_page_footer]
        for part in parts:
            yield from _iter_container(part)


def _iter_container(container):
    yield from getattr(container, "paragraphs", [])
    for table in getattr(container, "tables", []):
        yield from iter_table_paragraphs(table)


def iter_table_paragraphs(table):
    for row in table.rows:
        for tc in row._tr.tc_lst:
            from docx.table import _Cell
            cell = _Cell(tc, table)
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from iter_table_paragraphs(nested)


# run 안에서 텍스트가 아닌 내용물. 문단 텍스트를 갈아끼울 때 지워지면 안 된다.
# (사용인감 도장 이미지가 서명 문단의 run 안에 <w:drawing> 으로 들어 있다.)
NON_TEXT_TAGS = ("drawing", "pict", "object", "AlternateContent")


def _detach_non_text(runs):
    """run 들에서 그림/도형 요소를 떼어내 순서대로 돌려준다."""
    detached = []
    for run in runs:
        for child in list(run._element):
            if child.tag.split("}")[-1] in NON_TEXT_TAGS:
                run._element.remove(child)
                detached.append(child)
    return detached


def set_paragraph_text(paragraph, text):
    """첫 run 의 서식을 유지하면서 문단 텍스트를 교체한다.

    python-docx 의 run.text 설정은 run 안의 <w:drawing> 까지 날려버리므로,
    그림 요소는 미리 떼어놨다가 다시 붙인다.
    """
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    detached = _detach_non_text(runs)
    runs[0].text = text
    for run in runs[1:]:
        run._element.getparent().remove(run._element)
    for child in detached:
        runs[0]._element.append(child)


def replace_in_paragraph(paragraph, mapping):
    """문단 텍스트에 {{KEY}} 가 있으면 값으로 치환한다. 치환 시 True."""
    text = paragraph.text
    if "{{" not in text:
        return False
    new_text = text
    for key, value in mapping.items():
        new_text = new_text.replace("{{%s}}" % key, "" if value is None else str(value))
    if new_text == text:
        return False
    set_paragraph_text(paragraph, new_text)
    return True


def replace_everywhere(doc, mapping):
    for paragraph in iter_paragraphs(doc):
        replace_in_paragraph(paragraph, mapping)


def set_cell_text(cell, text):
    """셀 텍스트 교체. '\n' 은 셀 안의 별도 문단으로 나뉜다."""
    lines = str(text).split("\n")
    paragraphs = cell.paragraphs
    # 첫 문단을 서식 원본으로 삼아 부족한 만큼 복제한다.
    while len(paragraphs) < len(lines):
        new_p = copy.deepcopy(paragraphs[-1]._p)
        paragraphs[-1]._p.addnext(new_p)
        paragraphs = cell.paragraphs
    for idx, line in enumerate(lines):
        set_paragraph_text(paragraphs[idx], line)
    for extra in paragraphs[len(lines):]:
        extra._p.getparent().remove(extra._p)


def trim_spacers_after(table, max_remove):
    """표 바로 뒤에 있는 빈 문단을 최대 max_remove 개 지운다.

    내역 표에 줄이 늘어나면 서명부가 다음 장으로 밀리므로, 늘어난 줄 수만큼
    아래쪽 빈 줄을 걷어내 한 장에 유지한다. 글자나 그림이 있는 문단은 건드리지 않는다.
    제거한 개수를 돌려준다.
    """
    if max_remove <= 0:
        return 0
    removed = 0
    node = table._tbl.getnext()
    while node is not None and removed < max_remove:
        nxt = node.getnext()
        if node.tag.split("}")[-1] != "p":
            break
        from docx.text.paragraph import Paragraph
        paragraph = Paragraph(node, table._parent)
        has_content = paragraph.text.strip() or any(
            child.tag.split("}")[-1] in NON_TEXT_TAGS
            for run in paragraph.runs for child in run._element
        )
        if has_content:
            break
        node.getparent().remove(node)
        removed += 1
        node = nxt
    return removed


def clone_row(table, row_idx):
    """표의 특정 행을 그대로 복제해 바로 아래에 삽입하고 새 행을 돌려준다."""
    src_tr = table.rows[row_idx]._tr
    new_tr = copy.deepcopy(src_tr)
    src_tr.addnext(new_tr)
    return table.rows[row_idx + 1]


def delete_row(table, row_idx):
    tr = table.rows[row_idx]._tr
    tr.getparent().remove(tr)
